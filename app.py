import streamlit as st
import pandas as pd
import numpy as np
from scipy.interpolate import interp1d
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
import io
import re
import tempfile
import os
import json
import plotly.graph_objects as go
import plotly.io as pio
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

# ==================== 字体设置函数 ====================
def setup_matplotlib_font():
    """配置 matplotlib 中文字体，跨平台支持。
    优先尝试平台特定字体，若不可用则扫描系统字体中的 CJK 字体。
    """
    import platform
    # 先用 findSystemFonts 强制触发字体扫描，确保 apt 安装的字体被识别
    _ = fm.findSystemFonts()
    cjk_keywords = ['wqy', 'wenquan', 'noto.*cjk', 'noto.*sc', 'simhei', 'yahei',
                    'songti', 'heiti', 'droid.*fall', 'cjk']
    import re
    found_cjk = []
    for fpath in fm.findSystemFonts():
        fname = fpath.lower()
        if any(re.search(k, fname) for k in cjk_keywords):
            try:
                fp = fm.FontProperties(fname=fpath)
                found_cjk.append(fp.get_name())
            except Exception:
                continue
    # 平台默认候选字体
    if platform.system() == 'Windows':
        candidates = ['Microsoft YaHei', 'SimHei', 'DengXian']
    elif platform.system() == 'Linux':
        # WenQuanYi 通常名称为 "WenQuanYi Micro Hei"
        candidates = ['WenQuanYi Micro Hei', 'WenQuanYi Micro Hei Mono',
                      'Noto Sans CJK SC', 'Noto Sans SC',
                      'Droid Sans Fallback']
    else:
        candidates = ['STHeiti', 'Arial Unicode MS', 'PingFang SC']

    candidates = list(dict.fromkeys(candidates + found_cjk))
    plt.rcParams['font.sans-serif'] = candidates
    plt.rcParams['font.family'] = 'sans-serif'
    plt.rcParams['axes.unicode_minus'] = False

# ==================== 全局配置 ====================
st.set_page_config(
    page_title="拉伸测试报告生成器",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
div[data-testid="stColorPickerBlock"] {
    width: 82px !important; min-width: 82px !important;
    height: 12px !important; min-height: 12px !important;
    overflow: hidden;
}
div[data-testid="stColorPickerBlock"] > div:first-child {
    width: 82px !important; height: 12px !important;
}
div[data-testid="stColorPickerBlock"] input,
div[data-testid="stColorPickerBlock"] label {
    display: none !important;
}
div[data-testid="stSelectbox"] label p,
div[data-testid="stSelectbox"] div[data-baseweb="select"] span {
    font-size: 0.8rem !important;
}
</style>
""", unsafe_allow_html=True)

# ==================== 原始数据解析函数（通用） ====================
GROUP_STRIDE = 11
HEADER_ROW = 8
DATA_START_ROW = 10
HEADER_ROW_TESTID = 0
HEADER_ROW_DIMS = 2

COL_FORCE = 1      # 荷重(kgf)
COL_DISP = 2       # 位移(mm)
COL_STRESS = 9     # 应力(MPa)
COL_STRAIN = 10    # 应变(%)

def find_all_test_groups(df_raw):
    groups = []
    for col_idx in range(df_raw.shape[1]):
        cell_val = df_raw.iloc[HEADER_ROW_TESTID, col_idx]
        if pd.notna(cell_val) and "测试编号:" in str(cell_val):
            groups.append((col_idx, str(cell_val).strip()))
    return groups

def extract_test_dimensions(df_raw, group_offset):
    """提取一个测试组的全部维度信息（行1~行5），返回键值对字典。"""
    dim_info = {}

    # 第1行（索引1）：试品名称（第2列）
    specimen_name = df_raw.iloc[1, group_offset + 1] if group_offset + 1 < df_raw.shape[1] else None
    if pd.notna(specimen_name):
        dim_info["试品名称"] = str(specimen_name).strip()

    # 第2行（索引2）：长度, mm, 荷重, N, 形状, ..., 面积, value
    row2 = df_raw.iloc[2, group_offset:group_offset + GROUP_STRIDE]
    for i in range(0, len(row2), 2):
        key = row2.iloc[i]
        val = row2.iloc[i + 1] if i + 1 < len(row2) else None
        if pd.notna(key) and pd.notna(val):
            key_str = str(key).strip()
            if key_str == "长度":
                key_str = "长度单位"
            elif key_str == "荷重":
                key_str = "荷重单位"
            dim_info[key_str] = str(val).strip()

    # 第3行（索引3）：宽度{对边长}, value, 厚度, value, 倒角半径, value, ...
    row3 = df_raw.iloc[3, group_offset:group_offset + GROUP_STRIDE]
    for i in range(0, len(row3), 2):
        key = row3.iloc[i]
        val = row3.iloc[i + 1] if i + 1 < len(row3) else None
        if pd.notna(key) and pd.notna(val):
            dim_info[str(key).strip()] = str(val).strip()

    # 第4行（索引4）：标距, value, 长度, value, ...
    row4 = df_raw.iloc[4, group_offset:group_offset + GROUP_STRIDE]
    for i in range(0, len(row4), 2):
        key = row4.iloc[i]
        val = row4.iloc[i + 1] if i + 1 < len(row4) else None
        if pd.notna(key) and pd.notna(val):
            dim_info[str(key).strip()] = str(val).strip()

    # 兼容旧格式：若未提取到标距，尝试从旧位置读取
    if "标距" not in dim_info:
        gauge_raw = df_raw.iloc[HEADER_ROW_DIMS + 2, group_offset + 1]  # 原标距位置
        if pd.notna(gauge_raw):
            dim_info["标距"] = str(gauge_raw).strip()
    if "面积" not in dim_info:
        area_raw = df_raw.iloc[HEADER_ROW_DIMS, group_offset + 7]  # 原面积位置
        if pd.notna(area_raw):
            dim_info["面积"] = str(area_raw).strip()

    return dim_info

def read_columns_meta(df_raw, group_offset):
    headers = []
    for i in range(GROUP_STRIDE):
        col_idx = group_offset + i
        header = df_raw.iloc[HEADER_ROW, col_idx]
        unit = df_raw.iloc[HEADER_ROW + 1, col_idx]
        h = str(header).strip() if pd.notna(header) else ""
        u = str(unit).strip() if pd.notna(unit) else ""
        if h and u:
            col_name = f"{h}({u})"
        elif h:
            col_name = h
        elif u:
            col_name = f"({u})"
        else:
            col_name = f"列{col_idx}"
        headers.append(col_name)
    return headers

def extract_group_data(df_raw, group_offset):
    all_data = {}
    for i in range(GROUP_STRIDE):
        col_idx = group_offset + i
        data = pd.to_numeric(df_raw.iloc[DATA_START_ROW:, col_idx], errors='coerce').dropna().values
        all_data[i] = data
    force_kgf = all_data.get(COL_FORCE, [])
    disp = all_data.get(COL_DISP, [])
    stress = all_data.get(COL_STRESS, [])
    strain = all_data.get(COL_STRAIN, [])

    # 统一截断长度
    min_len = len(force_kgf)
    for arr in [disp, stress, strain]:
        if arr is not None:
            min_len = min(min_len, len(arr))
    if min_len == 0:
        return all_data, force_kgf, disp, None, None

    force_kgf = force_kgf[:min_len]
    disp = disp[:min_len]
    stress = stress[:min_len] if len(stress) >= min_len else None
    strain = strain[:min_len] if len(strain) >= min_len else None
    return all_data, force_kgf, disp, stress, strain

# ==================== 性能计算函数 ====================
def calculate_mechanical_properties(force_kgf, disp, stress, strain, gauge_length, area):
    if len(force_kgf) == 0 or stress is None or strain is None:
        return {k: 0.0 for k in ["max_force_N", "max_disp", "max_strain_pct",
                                 "tensile_strength", "E_modulus", "yield_stress", "yield_strain",
                                 "break_stress", "break_strain"]} | {"gauge_length": gauge_length, "area": area}
    max_force_N = np.max(force_kgf) * 9.80665
    idx_max = np.argmax(force_kgf)
    max_disp = disp[idx_max]
    max_strain_pct = max_disp / gauge_length * 100.0 if gauge_length > 0 else 0.0
    tensile_strength = np.max(stress)
    si = np.argsort(strain)
    ss = stress[si]
    sp = strain[si]
    sd = sp / 100.0
    break_stress = float(ss[-1])
    break_strain = float(sp[-1])
    mask = sd <= 0.01
    if np.sum(mask) < 3:
        mask = np.arange(len(sd)) < int(0.01 * len(sd))
    x = sd[mask]
    y = ss[mask]
    E_modulus = np.polyfit(x, y, 1)[0] if len(x) > 1 else 0.0
    if E_modulus > 0 and len(ss) > 1:
        off = 0.002
        off_line = E_modulus * (sd - off)
        diff = ss - off_line
        sc = np.where(np.diff(np.sign(diff)) != 0)[0]
        if len(sc) > 0:
            idx = sc[0]
            x1, x2 = sd[idx], sd[idx + 1]
            y1, y2 = diff[idx], diff[idx + 1]
            if y2 != y1:
                t = -y1 / (y2 - y1)
                ys_d = x1 + t * (x2 - x1)
                fi = interp1d(sd, ss, kind='linear', fill_value='extrapolate')
                yield_stress = float(fi(ys_d))
            else:
                yield_stress = float(ss[idx])
                ys_d = float(sd[idx])
        else:
            yield_stress = tensile_strength
            ys_d = float(sd[-1])
    else:
        yield_stress = tensile_strength
        ys_d = float(sd[-1]) if len(sd) > 0 else 0.0
    return {
        "max_force_N": max_force_N, "max_disp": max_disp,
        "max_strain_pct": max_strain_pct,
        "tensile_strength": tensile_strength, "E_modulus": E_modulus,
        "yield_stress": yield_stress, "yield_strain": ys_d * 100.0,
        "break_stress": break_stress, "break_strain": break_strain,
        "gauge_length": gauge_length, "area": area,
    }

def calculate_peel_properties(force_kgf, disp, width_mm):
    if len(force_kgf) == 0:
        return {"peel_strength_gf_cm": 0, "avg_force_N": 0, "max_force_N": 0, "min_force_N": 0}
    n = len(force_kgf)
    start = int(n * 0.2)
    end = int(n * 0.8)
    if end <= start:
        start = 0
        end = n
    stable_force_kgf = force_kgf[start:end]
    avg_force_kgf = np.mean(stable_force_kgf)
    max_force_kgf = np.max(force_kgf)
    min_force_kgf = np.min(force_kgf)
    width_cm = width_mm / 10.0
    peel_strength_gf_cm = avg_force_kgf * 1000 / width_cm if width_cm > 0 else 0
    return {
        "peel_strength_gf_cm": peel_strength_gf_cm,
        "avg_force_N": avg_force_kgf * 9.80665,
        "max_force_N": max_force_kgf * 9.80665,
        "min_force_N": min_force_kgf * 9.80665
    }

def parse_test_id(full_text):
    text = full_text
    if "测试编号:" in text:
        text = text.split("测试编号:", 1)[-1].strip()
    ts_match = re.search(r'\s+(\d{4}-\d{1,2}-\d{1,2}\s+\d{1,2}:\d{2}:\d{2}(?:\.\d+)?)\s*$', text)
    timestamp = ""
    clean_id = text
    if ts_match:
        timestamp = ts_match.group(1).strip()
        clean_id = text[:ts_match.start()].strip()
    batch_no = re.sub(r'-\d+\s*$', '', clean_id)
    return clean_id, batch_no, timestamp

def _find_col(df, *patterns):
    for p in patterns:
        for c in df.columns:
            if p in c:
                return c
    return None

# ==================== 编辑数据相关函数 ====================
def get_current_data_for_group(test_id):
    if test_id not in st.session_state.edited_data:
        return None
    df = st.session_state.edited_data[test_id]
    force_col = _find_col(df, "荷重", "力")
    disp_col = _find_col(df, "位移")
    stress_col = _find_col(df, "应力")
    strain_col = _find_col(df, "应变")
    if force_col is None or disp_col is None:
        return None
    return {
        "force": df[force_col].values,
        "disp": df[disp_col].values,
        "stress": df[stress_col].values if stress_col else None,
        "strain": df[strain_col].values if strain_col else None,
    }

def recalc_all_properties():
    new_props = []
    for test_id in st.session_state.test_ids:
        data = get_current_data_for_group(test_id)
        if data is None:
            continue
        dim_info = st.session_state.raw_group_data[test_id].get("dim_info", {})
        try:
            gauge = float(dim_info.get("标距", st.session_state.raw_group_data[test_id].get("gauge_length", 50)))
        except:
            gauge = 50.0
        try:
            area = float(dim_info.get("面积", st.session_state.raw_group_data[test_id].get("area", 0)))
        except:
            area = 0.0
        try:
            width = float(dim_info.get("宽度{对边长}", st.session_state.raw_group_data[test_id].get("width", 0)))
        except:
            width = 0.0

        if st.session_state.test_type == "拉伸性能测试":
            props = calculate_mechanical_properties(
                data["force"], data["disp"], data["stress"], data["strain"],
                gauge, area
            )
        else:
            props = calculate_peel_properties(data["force"], data["disp"], width)
        new_props.append(props)
    st.session_state.all_props = new_props
    return new_props

def reset_data_for_group(test_id):
    raw = st.session_state.raw_group_data[test_id]
    col_meta = raw.get("_col_meta", [])
    all_cols = raw.get("_all_data", {})
    target_len = len(raw["force"])
    df_dict = {}
    for idx, col_name in enumerate(col_meta):
        arr = all_cols.get(idx, [])
        if len(arr) >= target_len:
            df_dict[col_name] = arr[:target_len]
        elif len(arr) > 0:
            try:
                padded = np.pad(arr.astype(float), (0, target_len - len(arr)), constant_values=np.nan)
                df_dict[col_name] = padded
            except:
                padded = list(arr) + [""] * (target_len - len(arr))
                df_dict[col_name] = padded
        else:
            df_dict[col_name] = [np.nan] * target_len
    df = pd.DataFrame(df_dict).reset_index(drop=True)
    st.session_state.edited_data[test_id] = df
    recalc_all_properties()

# ==================== 导出/导入功能 ====================
def export_edited_data(test_ids, raw_group_data, all_props, edited_data):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        meta_data = []
        for idx, test_id in enumerate(test_ids):
            dim_info = raw_group_data[test_id].get("dim_info", {})
            props = all_props[idx] if all_props else {}
            row = {"测试组ID": test_id}
            for k, v in dim_info.items():
                try:
                    row[f"尺寸_{k}"] = float(v)
                except (ValueError, TypeError):
                    row[f"尺寸_{k}"] = str(v)
            if props:
                for pk, pv in props.items():
                    if isinstance(pv, (int, float)):
                        row[f"性能_{pk}"] = round(pv, 4)
                    else:
                        row[f"性能_{pk}"] = str(pv)
            meta_data.append(row)
        df_meta = pd.DataFrame(meta_data)
        df_meta.to_excel(writer, sheet_name="元数据", index=False)
        for test_id in test_ids:
            df = edited_data[test_id]
            df.to_excel(writer, sheet_name=test_id, index=False)
    output.seek(0)
    return output

def import_edited_data(uploaded_file):
    try:
        file_bytes = uploaded_file.getvalue()
        file_buffer = io.BytesIO(file_bytes)
        xlsx = pd.ExcelFile(file_buffer)
        if "元数据" not in xlsx.sheet_names:
            st.error('无效文件：缺少"元数据"工作表')
            return False
        df_meta = pd.read_excel(io.BytesIO(file_bytes), sheet_name="元数据")
        required_cols = ["测试组ID", "标距(mm)", "面积(mm²)"]
        if not all(col in df_meta.columns for col in required_cols):
            st.error("元数据工作表缺少必要列")
            return False

        test_ids = []
        raw_group_data = {}
        edited_data = {}

        for _, row in df_meta.iterrows():
            tid = row["测试组ID"]
            gauge = row["标距(mm)"]
            area = row["面积(mm²)"]
            width = row.get("宽度(mm)", 0)
            if tid not in xlsx.sheet_names:
                st.warning(f"测试组 {tid} 没有数据工作表，跳过")
                continue
            df_data = pd.read_excel(io.BytesIO(file_bytes), sheet_name=tid)
            if df_data.empty:
                continue
            force_col = _find_col(df_data, "荷重", "力")
            disp_col = _find_col(df_data, "位移")
            stress_col = _find_col(df_data, "应力")
            strain_col = _find_col(df_data, "应变")
            if force_col is None or disp_col is None:
                st.warning(f"测试组 {tid} 找不到荷重/位移列，跳过")
                continue
            dim_info = {"标距": str(gauge), "面积": str(area), "宽度{对边长}": str(width)}
            raw_group_data[tid] = {
                "force": df_data[force_col].values,
                "disp": df_data[disp_col].values,
                "stress": df_data[stress_col].values if stress_col else None,
                "strain": df_data[strain_col].values if strain_col else None,
                "gauge_length": gauge,
                "area": area,
                "width": width,
                "thickness": 0.0,
                "_col_meta": list(df_data.columns),
                "_all_data": {idx: df_data[c].values for idx, c in enumerate(df_data.columns)},
                "dim_info": dim_info,
            }
            edited_data[tid] = df_data.copy()
            test_ids.append(tid)

        if len(test_ids) == 0:
            st.error("未找到任何有效的测试组数据")
            return False

        st.session_state.test_ids = test_ids
        st.session_state.raw_group_data = raw_group_data
        st.session_state.edited_data = edited_data
        st.session_state.all_props = []  # 修复：先置空，后续手动计算
        recalc_all_properties()  # 重新计算所有性能指标
        first_id = test_ids[0]
        batch_no_match = re.sub(r'-\d+$', '', first_id)
        st.session_state.batch_no = batch_no_match
        st.session_state.timestamp = "导入数据"
        return True
    except Exception as e:
        st.error(f"导入失败: {str(e)}")
        return False

# ==================== 绘图函数（Plotly） ====================
def plot_plotly_chart(selected_groups, x_var_key, y_var_key, filter_method,
                      strain_min, strain_max, row_start, row_end,
                      x_label, y_label, line_width, color_mode, custom_colors,
                      edited_data):
    strain_col = None
    if filter_method == "按应变范围" and selected_groups:
        for c in edited_data[selected_groups[0]].columns:
            if "应变" in c:
                strain_col = c
                break

    palette = ['#1a237e', '#ef5350', '#2e7d32', '#ff8f00', '#6a1b9a',
               '#00838f', '#d81b60', '#3e2723', '#558b2f', '#01579b']

    fig = go.Figure()
    fig.update_layout(
        title_text=f"{y_label} - {x_label} 曲线",
        xaxis_title=x_label,
        yaxis_title=y_label,
        hovermode='closest',
        plot_bgcolor='#fafafa',
        paper_bgcolor='white',
        font=dict(family='Microsoft YaHei, Segoe UI, Arial', size=12, color='#333333'),
        title_font=dict(size=14, color='#1a237e', family='Microsoft YaHei, Segoe UI'),
        legend=dict(
            yanchor="middle", y=0.5,
            xanchor="left", x=1.02,
            bgcolor='rgba(255,255,255,0.85)',
            bordercolor='#dddddd', borderwidth=1,
            font=dict(size=10)
        ),
        margin=dict(l=50, r=250, t=50, b=50),
        xaxis=dict(gridcolor='#e0e0e0', gridwidth=0.5, zeroline=False, linecolor='#cccccc', linewidth=0.8),
        yaxis=dict(gridcolor='#e0e0e0', gridwidth=0.5, zeroline=False, linecolor='#cccccc', linewidth=0.8),
        width=None, height=400,
    )

    for idx, test_id in enumerate(selected_groups):
        df = edited_data[test_id]
        if x_var_key not in df.columns or y_var_key not in df.columns:
            continue
        x_raw = df[x_var_key].values
        y_raw = df[y_var_key].values
        if len(x_raw) == 0 or len(y_raw) == 0:
            continue

        if filter_method == "按应变范围" and strain_col and strain_col in df.columns:
            strain_arr = df[strain_col].values
            mask = (strain_arr >= strain_min) & (strain_arr <= strain_max)
            x_plot = x_raw[mask].tolist()
            y_plot = y_raw[mask].tolist()
        elif filter_method == "按行号范围":
            start = max(0, row_start)
            end = min(len(x_raw) - 1, row_end)
            x_plot = x_raw[start:end + 1].tolist()
            y_plot = y_raw[start:end + 1].tolist()
        else:
            x_plot = x_raw.tolist()
            y_plot = y_raw.tolist()

        if len(x_plot) == 0:
            continue

        if color_mode == "自定义每个测试组" and test_id in custom_colors:
            color = custom_colors[test_id]
        else:
            color = palette[idx % len(palette)]

        fig.add_trace(go.Scatter(
            x=x_plot, y=y_plot,
            mode='markers',
            name=test_id,
            line=dict(width=line_width * 0.5, color=color),
            marker=dict(size=line_width, color=color),
            hovertemplate=f'{x_label}: %{{x}}<br>{y_label}: %{{y}}<extra>{test_id}</extra>'
        ))

    if len(fig.data) > 0:
        all_y = np.concatenate([np.array(t.y) for t in fig.data if t.y is not None])
        if len(all_y) > 0 and np.all(all_y >= 0):
            fig.update_yaxes(range=[0, max(all_y) * 1.05])
        all_x = np.concatenate([np.array(t.x) for t in fig.data if t.x is not None])
        if len(all_x) > 0 and np.all(all_x >= 0):
            fig.update_xaxes(range=[0, max(all_x) * 1.05])

    return fig

# ==================== Word 报告生成 ====================
def _set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def _set_cell_font(cell, size=Pt(10), bold=False, font_name='Microsoft YaHei', alignment=WD_ALIGN_PARAGRAPH.CENTER):
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        for run in paragraph.runs:
            run.font.size = size
            run.bold = bold
            run.font.name = font_name
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:rFont'), font_name)
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rPr.append(rFonts)
            else:
                rFonts.set(qn('w:eastAsia'), font_name)
    for paragraph in cell.paragraphs:
        if not paragraph.runs:
            run = paragraph.add_run(cell.text)
            run.font.size = size
            run.bold = bold
            run.font.name = font_name

def _set_table_col_widths(table, widths_cm):
    for row in table.rows:
        for i, width in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(width)

# ==================== 统一 Matplotlib 绘图函数（参考 app_0.py） ====================
def plot_matplotlib_chart(selected_groups, edited_data, x_var, y_var,
                          x_label, y_label, line_width, custom_colors):
    """
    使用 matplotlib 绘制曲线图，返回 figure 对象。
    供 Word 报告生成使用，与 app_0.py 的 plot_custom_chart 结构一致。
    """
    setup_matplotlib_font()
    fig, ax = plt.subplots(figsize=(12, 5))
    palette = ['#1a237e', '#ef5350', '#2e7d32', '#ff8f00', '#6a1b9a',
               '#00838f', '#d81b60', '#3e2723', '#558b2f', '#01579b']
    plotted = 0
    for idx, tid in enumerate(selected_groups):
        if tid not in edited_data:
            continue
        df = edited_data[tid]
        if x_var not in df.columns or y_var not in df.columns:
            continue
        xv = df[x_var].values
        yv = df[y_var].values
        if len(xv) == 0:
            continue
        color = custom_colors.get(tid, palette[idx % len(palette)])
        ax.plot(xv, yv, color=color, linewidth=max(0.5, line_width * 0.5),
                marker='o', markersize=line_width, label=tid)
        plotted += 1
    if plotted == 0:
        ax.text(0.5, 0.5, '无有效数据可绘图', transform=ax.transAxes,
                ha='center', va='center', fontsize=14, color='gray')
    ax.set_xlabel(x_label, fontsize=11)
    ax.set_ylabel(y_label, fontsize=11)
    ax.set_title(f"{y_label} - {x_label} 曲线", fontsize=13, color='#1a237e')
    ax.grid(True, alpha=0.3)
    if plotted > 0:
        ax.legend(fontsize=9)
    fig.tight_layout()
    return fig


def generate_word_report_bytes(selected_groups, x_var, y_var, x_label, y_label,
                               line_width, color_mode, custom_colors_tuple,
                               company_name, report_title, batch_no, timestamp,
                               tester_name, client_name, test_standard,
                               test_type, all_props, raw_group_data, test_ids,
                               edited_data):
    """
    所有数据通过参数传入，不依赖 st.session_state。
    """
    custom_colors = dict(custom_colors_tuple)
    chart_buffer = None
    try:
        fig = plot_matplotlib_chart(
            selected_groups, edited_data, x_var, y_var,
            x_label, y_label, line_width, custom_colors
        )
        chart_buffer = io.BytesIO()
        fig.savefig(chart_buffer, format='png', dpi=150, facecolor='white')
        plt.close(fig)
    except Exception:
        chart_buffer = None

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(company_name)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Microsoft YaHei'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{company_name}{report_title}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)

    header_table = doc.add_table(rows=5, cols=2)
    header_table.style = 'Table Grid'
    header_data = [("测试批号:", batch_no), ("测试人员:", tester_name),
                   ("客户名称:", client_name), ("测试标准:", test_standard),
                   ("测试日期:", timestamp)]
    for i, (label, val) in enumerate(header_data):
        for j, txt in enumerate([label, val]):
            c = header_table.rows[i].cells[j]
            c.text = ""
            rr = c.paragraphs[0].add_run(txt)
            rr.font.size = Pt(11)
            rr.font.name = 'Microsoft YaHei'
            if j == 0:
                rr.bold = True
                _set_cell_shading(c, "e8eaf6")
    _set_table_col_widths(header_table, [3.0, 23.2])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("测试结果：")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'

    USABLE_WIDTH = 26.2
    if test_type == "拉伸性能测试":
        headers = ["测试编号", "最大荷重(N)", "最大荷重位移(mm)", "最大荷重伸长率(%)",
                   "抗拉强度(MPa)", "弹性模量(Ei)(MPa)", "屈服强度(MPa)", "屈服伸长率(%)",
                   "断裂强度(MPa)", "断裂伸长率(%)", "标距(mm)", "面积(mm²)"]
        col_widths = [3.2] + [2.0]*9 + [1.8, 1.8]
        col_widths = [w * USABLE_WIDTH / sum(col_widths) for w in col_widths]

        num_data_rows = len(all_props)
        total_rows = 1 + num_data_rows + 3
        table = doc.add_table(rows=total_rows, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            _set_cell_shading(cell, "1a237e")
            _set_cell_font(cell, size=Pt(9), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_table_col_widths(table, col_widths)

        for idx, (props, test_id) in enumerate(zip(all_props, test_ids)):
            row = table.rows[1 + idx]
            values = [
                test_id,
                f"{props['max_force_N']:.3f}", f"{props['max_disp']:.3f}",
                f"{props['max_strain_pct']:.3f}", f"{props['tensile_strength']:.3f}",
                f"{props['E_modulus']:.3f}", f"{props['yield_stress']:.3f}",
                f"{props['yield_strain']:.3f}", f"{props['break_stress']:.3f}",
                f"{props['break_strain']:.3f}", f"{props['gauge_length']:.3f}",
                f"{props['area']:.3f}",
            ]
            for j, val in enumerate(values):
                row.cells[j].text = val
                _set_cell_font(row.cells[j], size=Pt(9), bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            if idx % 2 == 1:
                for j in range(len(headers)):
                    _set_cell_shading(row.cells[j], "f5f5ff")

        stat_rows = [("最大值 Max", np.max), ("最小值 Min", np.min), ("平均值 X-bar", np.mean)]
        stat_fields = ["max_force_N", "max_disp", "max_strain_pct",
                       "tensile_strength", "E_modulus", "yield_stress",
                       "yield_strain", "break_stress", "break_strain"]
        stat_colors = ["e3f2fd", "fce4ec", "e8f5e9"]
        for s_idx, (label, func) in enumerate(stat_rows):
            row = table.rows[1 + num_data_rows + s_idx]
            row.cells[0].text = label
            _set_cell_shading(row.cells[0], stat_colors[s_idx])
            _set_cell_font(row.cells[0], size=Pt(9), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for f_idx, field in enumerate(stat_fields):
                vals = [p[field] for p in all_props]
                stat_val = func(vals)
                row.cells[1 + f_idx].text = f"{stat_val:.3f}"
                _set_cell_shading(row.cells[1 + f_idx], stat_colors[s_idx])
                _set_cell_font(row.cells[1 + f_idx], size=Pt(9), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            gauge_vals = [p["gauge_length"] for p in all_props]
            row.cells[len(headers) - 2].text = f"{func(gauge_vals):.3f}"
            _set_cell_shading(row.cells[len(headers) - 2], stat_colors[s_idx])
            _set_cell_font(row.cells[len(headers) - 2], size=Pt(9), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            area_vals = [p["area"] for p in all_props]
            row.cells[len(headers) - 1].text = f"{func(area_vals):.3f}"
            _set_cell_shading(row.cells[len(headers) - 1], stat_colors[s_idx])
            _set_cell_font(row.cells[len(headers) - 1], size=Pt(9), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    else:
        # 剥离测试
        headers = ["测试编号", "180°剥离平均强度(gf/cm)", "宽度(cm)", "单区间最大荷重(N)",
                   "单区间最小荷重(N)", "单区间荷重平均值(N)"]
        col_widths = [4.0, 5.0, 2.5, 4.5, 4.5, 4.5]
        col_widths = [w * USABLE_WIDTH / sum(col_widths) for w in col_widths]

        num_data_rows = len(all_props)
        total_rows = 1 + num_data_rows + 3
        table = doc.add_table(rows=total_rows, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            _set_cell_shading(cell, "1a237e")
            _set_cell_font(cell, size=Pt(10), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_table_col_widths(table, col_widths)

        for idx, (props, test_id) in enumerate(zip(all_props, test_ids)):
            width_cm = raw_group_data[test_id].get("width", 0) / 10.0
            row = table.rows[1 + idx]
            values = [
                test_id,
                f"{props['peel_strength_gf_cm']:.3f}",
                f"{width_cm:.3f}",
                f"{props['max_force_N']:.3f}",
                f"{props['min_force_N']:.3f}",
                f"{props['avg_force_N']:.3f}"
            ]
            for j, val in enumerate(values):
                row.cells[j].text = val
                _set_cell_font(row.cells[j], size=Pt(10), bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            if idx % 2 == 1:
                for j in range(len(headers)):
                    _set_cell_shading(row.cells[j], "f5f5ff")

        stat_rows = [("最大值 Max", np.max), ("最小值 Min", np.min), ("平均值 X-bar", np.mean)]
        stat_fields = ["peel_strength_gf_cm", "max_force_N", "min_force_N", "avg_force_N"]
        stat_colors = ["e3f2fd", "fce4ec", "e8f5e9"]
        for s_idx, (label, func) in enumerate(stat_rows):
            row = table.rows[1 + num_data_rows + s_idx]
            row.cells[0].text = label
            _set_cell_shading(row.cells[0], stat_colors[s_idx])
            _set_cell_font(row.cells[0], size=Pt(10), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for f_idx, field in enumerate(stat_fields):
                vals = [p[field] for p in all_props]
                stat_val = func(vals)
                row.cells[1 + f_idx].text = f"{stat_val:.3f}"
                _set_cell_shading(row.cells[1 + f_idx], stat_colors[s_idx])
                _set_cell_font(row.cells[1 + f_idx], size=Pt(10), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            width_vals = [raw_group_data[t].get("width", 0) / 10.0 for t in test_ids]
            row.cells[len(headers) - 2].text = f"{func(width_vals):.3f}"
            _set_cell_shading(row.cells[len(headers) - 2], stat_colors[s_idx])
            _set_cell_font(row.cells[len(headers) - 2], size=Pt(10), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("测试曲线：")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'
    if chart_buffer is not None:
        chart_buffer.seek(0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart_buffer, width=Cm(24))
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(f"图表说明：X轴为{x_label}，Y轴为{y_label}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer.getvalue()

# ==================== HTML 报告生成 ====================
def generate_html_report(test_ids, all_props, batch_no, timestamp,
                         x_var_key, y_var_key, x_label, y_label, filter_method,
                         strain_min, strain_max, row_start, row_end,
                         selected_groups, line_width, custom_colors,
                         tester_name, client_name, test_standard,
                         test_type, edited_data, raw_group_data,
                         company_name, report_title):
    """
    所有数据通过参数传入，不依赖 st.session_state。
    """
    series_data = []
    raw_series_data = []
    all_col_names = []
    for test_id in test_ids:
        df = edited_data[test_id]
        if all_col_names == []:
            all_col_names = list(df.columns)
        if x_var_key not in df.columns or y_var_key not in df.columns:
            continue
        entry = {"name": test_id}
        for col in df.columns:
            entry[col] = df[col].values.tolist()
        raw_series_data.append(entry)

    if not raw_series_data:
        return "<html><body><h2>无可用数据</h2></body></html>"

    if test_type == "拉伸性能测试":
        headers = ["测试编号", "最大荷重(N)", "最大荷重位移(mm)", "最大荷重伸长率(%)",
                   "抗拉强度(MPa)", "弹性模量(Ei)(MPa)", "屈服强度(MPa)", "屈服伸长率(%)",
                   "断裂强度(MPa)", "断裂伸长率(%)", "标距(mm)", "面积(mm²)"]
        table_rows = []
        for idx, test_id in enumerate(test_ids):
            props = all_props[idx]
            row = [
                test_id,
                f"{props['max_force_N']:.3f}", f"{props['max_disp']:.3f}",
                f"{props['max_strain_pct']:.3f}", f"{props['tensile_strength']:.3f}",
                f"{props['E_modulus']:.3f}", f"{props['yield_stress']:.3f}",
                f"{props['yield_strain']:.3f}", f"{props['break_stress']:.3f}",
                f"{props['break_strain']:.3f}", f"{props['gauge_length']:.3f}",
                f"{props['area']:.3f}"
            ]
            table_rows.append(row)
        stat_fields = ["max_force_N", "max_disp", "max_strain_pct",
                       "tensile_strength", "E_modulus", "yield_stress",
                       "yield_strain", "break_stress", "break_strain"]
        stat_rows = []
        for label, func in [("最大值", np.max), ("最小值", np.min), ("平均值", np.mean)]:
            row = [label]
            for field in stat_fields:
                vals = [p[field] for p in all_props]
                row.append(f"{func(vals):.3f}")
            gauge_vals = [p["gauge_length"] for p in all_props]
            row.append(f"{func(gauge_vals):.3f}")
            area_vals = [p["area"] for p in all_props]
            row.append(f"{func(area_vals):.3f}")
            stat_rows.append(row)
    else:
        headers = ["测试编号", "180°剥离平均强度(gf/cm)", "宽度(cm)", "单区间最大荷重(N)",
                   "单区间最小荷重(N)", "单区间荷重平均值(N)"]
        table_rows = []
        for idx, test_id in enumerate(test_ids):
            props = all_props[idx]
            width_cm = raw_group_data[test_id].get("width", 0) / 10.0
            row = [
                test_id,
                f"{props['peel_strength_gf_cm']:.3f}",
                f"{width_cm:.3f}",
                f"{props['max_force_N']:.3f}",
                f"{props['min_force_N']:.3f}",
                f"{props['avg_force_N']:.3f}"
            ]
            table_rows.append(row)
        stat_fields = ["peel_strength_gf_cm", "max_force_N", "min_force_N", "avg_force_N"]
        stat_rows = []
        for label, func in [("最大值", np.max), ("最小值", np.min), ("平均值", np.mean)]:
            row = [label]
            for field in stat_fields:
                vals = [p[field] for p in all_props]
                row.append(f"{func(vals):.3f}")
            width_vals = [raw_group_data[t].get("width", 0) / 10.0 for t in test_ids]
            row.append(f"{func(width_vals):.3f}")
            stat_rows.append(row)

    default_colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf']
    color_map = {}
    for i, tid in enumerate(test_ids):
        if tid in custom_colors:
            color_map[tid] = custom_colors[tid]
        else:
            color_map[tid] = default_colors[i % len(default_colors)]


    html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>测试报告 - {batch_no}</title>
    <script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
    <style>
        body {{ font-family: 'Segoe UI', 'Microsoft YaHei', Arial, sans-serif; margin: 20px; background: white; }}
        .container {{ max-width: 1200px; margin: auto; background: white; padding: 24px 28px; border-radius: 12px; box-shadow: 0 2px 20px rgba(0,0,0,0.08); }}
        h1 {{ color: #1a237e; text-align: center; font-size: 1.6rem; margin-bottom: 0.2rem; }}
        h2 {{ color: #283593; text-align: center; font-size: 1.2rem; font-weight: 500; margin-top: 0.2rem; margin-bottom: 1.5rem; }}
        .info {{ background: linear-gradient(135deg, #e8eaf6, #c5cae9); padding: 12px 18px; border-radius: 8px; margin-bottom: 24px; border-left: 4px solid #1a237e; font-size: 0.95rem; }}
        table {{ border-collapse: collapse; width: 100%; margin: 16px 0; font-size: 0.9rem; }}
        th, td {{ border: 1px solid #d0d0d0; padding: 8px 10px; text-align: center; }}
        th {{ background-color: #1a237e; color: white; font-weight: 600; }}
        tr:nth-child(even) {{ background-color: #f8f8ff; }}
        .stat-row {{ background-color: #e8f5e9 !important; font-weight: bold; }}
        .controls {{ background: #f8f9ff; padding: 16px 20px; border-radius: 8px; margin-bottom: 20px; border: 1px solid #e0e0f0; }}
        .control-group {{ display: inline-block; margin-right: 20px; margin-bottom: 10px; }}
        label {{ font-weight: 600; margin-right: 5px; color: #333; }}
        select, input {{ padding: 6px 10px; border-radius: 6px; border: 1px solid #ccc; font-size: 0.9rem; }}
        select:focus, input:focus {{ outline: none; border-color: #3949ab; box-shadow: 0 0 0 2px rgba(57,73,171,0.15); }}
        button {{ background: linear-gradient(135deg, #3949ab, #1a237e); color: white; border: none; padding: 8px 20px; border-radius: 6px; cursor: pointer; font-weight: 500; transition: all 0.2s; }}
        button:hover {{ background: linear-gradient(135deg, #5c6bc0, #3949ab); box-shadow: 0 2px 8px rgba(26,35,126,0.3); transform: translateY(-1px); }}
        .color-input {{ width: 60px; height: 32px; cursor: pointer; }}
        .note {{ font-size: 0.85rem; color: #666; margin-top: 20px; padding: 12px; background: #fafafa; border-radius: 6px; border-left: 3px solid #3949ab; }}
        hr {{ margin: 20px 0; border: none; height: 1px; background: linear-gradient(90deg, transparent, #c5cae9, transparent); }}
        .chart-container {{ background: white; border-radius: 8px; border: 1px solid #eee; overflow: hidden; }}
        .section-title {{ font-size: 1.1rem; font-weight: 600; color: #1a237e; margin: 1.2rem 0 0.6rem 0; padding-bottom: 0.3rem; border-bottom: 2px solid #e8eaf6; }}
    </style>
</head>
<body>
<div class="container">
    <h1>{company_name}</h1>
    <h2>{company_name}{report_title}</h2>
    <table style="width:100%; border-collapse:collapse; margin-bottom:20px;">
        <tr>
            <td style="padding:4px 10px; border:1px solid #d0d0d0; background:#e8eaf6; font-weight:600; width:15%;">测试批号:</td>
            <td style="padding:4px 10px; border:1px solid #d0d0d0; width:35%;">{batch_no}</td>
            <td style="padding:4px 10px; border:1px solid #d0d0d0; background:#e8eaf6; font-weight:600; width:15%;">测试人员:</td>
            <td style="padding:4px 10px; border:1px solid #d0d0d0; width:35%;">{tester_name}</td>
        </tr>
        <tr>
            <td style="padding:4px 10px; border:1px solid #d0d0d0; background:#e8eaf6; font-weight:600;">客户名称:</td>
            <td style="padding:4px 10px; border:1px solid #d0d0d0;">{client_name}</td>
            <td style="padding:4px 10px; border:1px solid #d0d0d0; background:#e8eaf6; font-weight:600;">测试标准:</td>
            <td style="padding:4px 10px; border:1px solid #d0d0d0;">{test_standard}</td>
        </tr>
        <tr>
            <td style="padding:4px 10px; border:1px solid #d0d0d0; background:#e8eaf6; font-weight:600;">测试日期:</td>
            <td style="padding:4px 10px; border:1px solid #d0d0d0;" colspan="3">{timestamp}</td>
        </tr>
    </table>
    <h2 class="section-title">测试结果</h2>
    <table>
        <thead><tr>{''.join(f'<th>{h}</th>' for h in headers)}</thead>
        <tbody>{''.join('<tr>' + ''.join(f'<td>{v}</td>' for v in row) + '</tr>' for row in table_rows)}
        {''.join('<tr>' + ''.join(f'<td><b>{v}</b></td>' for v in row) + '</tr>' for row in stat_rows)}</tbody>
    </table>
    <h2 class="section-title">可交互曲线图</h2>
    <div class="controls">
        <div class="control-group"><label>X轴变量:</label><select id="xVar">
            {''.join(f'<option value="{c}" {"selected" if c == x_var_key else ""}>{c}</option>' for c in all_col_names)}
        </select></div>
        <div class="control-group"><label>Y轴变量:</label><select id="yVar">
            {''.join(f'<option value="{c}" {"selected" if c == y_var_key else ""}>{c}</option>' for c in all_col_names)}
        </select></div>
        <div class="control-group"><label>X轴标题:</label><input type="text" id="xLabel" value="{x_label}"></div>
        <div class="control-group"><label>Y轴标题:</label><input type="text" id="yLabel" value="{y_label}"></div>
        <div class="control-group"><label>线条粗细:</label><input type="range" id="lineWidth" min="0.5" max="5" step="0.1" value="{line_width}"></div>
        <div class="control-group"><label>测试组:</label><select id="groupSelect" multiple size="3">
            {''.join(f'<option value="{g}" {"selected" if g in selected_groups else ""}>{g}</option>' for g in test_ids)}
        </select><br><small>按住Ctrl多选</small></div>
        <div><button id="updateBtn">更新图表</button></div>
    </div>
    <div id="plotlyChart" class="chart-container"></div>
    <div class="controls" style="margin-top:10px">
        <h4 style="margin:0 0 10px 0; color:#1a237e;">自定义每组颜色</h4>
        {''.join(f'<div class="control-group"><label>{tid}:</label><input type="color" id="color_{tid}" value="{color_map[tid]}" class="color-input"></div>' for tid in test_ids)}
        <div><button id="applyColorsBtn">应用颜色</button></div>
    </div>
    <div class="note">
        <strong>操作提示：</strong> 图表支持缩放（滚轮）、平移（拖拽）、下载为PNG（右上角相机图标）；可任意切换X/Y变量、选择测试组、调整线条粗细和颜色。
    </div>
</div>
<script>
    const allRawData = {json.dumps(raw_series_data, ensure_ascii=False)};
    const allTestIds = {json.dumps(test_ids, ensure_ascii=False)};
    let currentColors = {json.dumps(color_map, ensure_ascii=False)};
    function applyFilterAndPlot() {{
        const xVar = document.getElementById('xVar').value;
        const yVar = document.getElementById('yVar').value;
        const xLabel = document.getElementById('xLabel').value;
        const yLabel = document.getElementById('yLabel').value;
        let selected = Array.from(document.getElementById('groupSelect').selectedOptions).map(opt => opt.value);
        selected.sort((a,b) => allTestIds.indexOf(a) - allTestIds.indexOf(b));
        const lineWidth = parseFloat(document.getElementById('lineWidth').value);
        const traces = [];
        for (let testId of selected) {{
            const group = allRawData.find(g => g.name === testId);
            if (!group) continue;
            const xPlot = group[xVar] || [];
            const yPlot = group[yVar] || [];
            if (xPlot.length === 0 || yPlot.length === 0) continue;
            const color = currentColors[testId] || '#1f77b4';
            traces.push({{
                x: xPlot,
                y: yPlot,
                mode: 'lines+markers',
                name: testId,
                line: {{ width: lineWidth * 0.5, color: color }},
                marker: {{ size: lineWidth, color: color }}
            }});
        }}
        const layout = {{
            title: `${{yLabel}} - ${{xLabel}} 曲线`,
            xaxis: {{ title: xLabel }},
            yaxis: {{ title: yLabel }},
            hovermode: 'closest',
            autosize: true,
            legend: {{ x: 1.02, y: 0.5, xanchor: "left", yanchor: "middle" }},
            margin: {{ l: 50, r: 280, t: 60, b: 50 }},
        }};
        Plotly.newPlot('plotlyChart', traces, layout, {{ responsive: true, displayModeBar: true }});
    }}
    function updateColors() {{
        const testIds = {json.dumps(test_ids, ensure_ascii=False)};
        for (let tid of testIds) {{
            const picker = document.getElementById('color_' + tid);
            if (picker) {{
                currentColors[tid] = picker.value;
            }}
        }}
        applyFilterAndPlot();
    }}
    window.addEventListener('resize', () => {{
        const container = document.getElementById('plotlyChart');
        if (container) Plotly.Plots.resize(container);
    }});
    document.getElementById('updateBtn').addEventListener('click', applyFilterAndPlot);
    document.getElementById('applyColorsBtn').addEventListener('click', updateColors);
    document.getElementById('lineWidth').addEventListener('input', function() {{ applyFilterAndPlot(); }});
    window.onload = applyFilterAndPlot;
</script>
</body>
</html>"""
    return html_template

# ==================== CSS 样式 ====================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Segoe UI', 'Microsoft YaHei', -apple-system, sans-serif; }
    .stApp { background: white; }
    .main > div { padding: 0 1rem; }
    .block-container { max-width: 1100px; padding: 1.2rem 1.5rem !important; margin: 0 auto; }
    .main-header { font-size: 1.3rem; font-weight: 700; color: #1a237e; text-align: center; margin-bottom: 0.1rem; margin-top: 1.3rem; letter-spacing: 0.3px; padding: 0.3rem 0 0 0; line-height: 1.3; }
    .sub-header { text-align: center; color: #5c6bc0; font-size: 0.82rem; margin-top: 0.3rem; margin-bottom: 0.9rem; font-weight: 400; }
    div[data-testid="stMetric"] { background: linear-gradient(135deg, #e8eaf6, #c5cae9); border-radius: 6px; padding: 10px 14px; border-left: 3px solid #1a237e; }
    div[data-testid="stMetric"] label { font-size: 0.72rem !important; font-weight: 600 !important; color: #1a237e !important; text-transform: uppercase; letter-spacing: 0.2px; }
    div[data-testid="stMetric"] div { font-size: 1.05rem !important; font-weight: 700 !important; color: #1a237e !important; }
    .stButton button, .stDownloadButton button { background: linear-gradient(135deg, #3949ab, #1a237e); color: white !important; border: none; border-radius: 6px; font-weight: 500; padding: 0.3rem 1rem; font-size: 0.82rem; transition: all 0.2s; }
    .stButton button:hover, .stDownloadButton button:hover { background: linear-gradient(135deg, #5c6bc0, #3949ab); box-shadow: 0 2px 8px rgba(26,35,126,0.3); transform: translateY(-1px); }
    .stDataFrame { border-radius: 6px; overflow: hidden; border: 1px solid #d0d0d0; font-size: 0.85rem; }
    .stDataFrame thead tr th { background-color: #1a237e !important; color: white !important; font-weight: 600; font-size: 0.82rem; }
    section[data-testid="stSidebar"] { background: #f0f2f5; border-right: 1px solid #e8e8e8; }
    .streamlit-expanderHeader { font-weight: 600; color: #1a237e; background: #f8f9ff; border-radius: 6px; padding: 0.3rem 0.7rem; font-size: 0.82rem; border: 1px solid #e0e0f0; }
    hr { margin: 0.6rem 0; border: none; height: 1px; background: linear-gradient(90deg, transparent, #c5cae9, transparent); }
    div[data-testid="stMarkdownContainer"] h3 { font-size: 1rem !important; font-weight: 600 !important; color: #1a237e !important; margin-top: 0.6rem !important; margin-bottom: 0.2rem !important; }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header">📊 材料测试数据交互式报告生成器</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">上传Excel测试数据 · 自由配置图表 · 一键生成专业报告</div>', unsafe_allow_html=True)

# ==================== 初始化 session_state ====================
if "edited_data" not in st.session_state:
    st.session_state.edited_data = {}
if "all_props" not in st.session_state:
    st.session_state.all_props = None
if "raw_group_data" not in st.session_state:
    st.session_state.raw_group_data = None
if "test_ids" not in st.session_state:
    st.session_state.test_ids = []
if "batch_no" not in st.session_state:
    st.session_state.batch_no = ""
if "timestamp" not in st.session_state:
    st.session_state.timestamp = ""
if "test_type" not in st.session_state:
    st.session_state.test_type = "拉伸性能测试"
if "company_name" not in st.session_state:
    st.session_state.company_name = "XXXX有限公司"
if "report_title" not in st.session_state:
    st.session_state.report_title = "测试报告"
if "selected_groups" not in st.session_state:
    st.session_state.selected_groups = set()
if "last_ids" not in st.session_state:
    st.session_state.last_ids = []
if "x_var" not in st.session_state:
    st.session_state.x_var = ""
if "y_var" not in st.session_state:
    st.session_state.y_var = ""
if "x_label" not in st.session_state:
    st.session_state.x_label = ""
if "y_label" not in st.session_state:
    st.session_state.y_label = ""
if "filter_method" not in st.session_state:
    st.session_state.filter_method = "无筛选"
if "line_width" not in st.session_state:
    st.session_state.line_width = 2
if "color_mode" not in st.session_state:
    st.session_state.color_mode = "自动分配"
if "tester_name" not in st.session_state:
    st.session_state.tester_name = ""
if "client_name" not in st.session_state:
    st.session_state.client_name = ""
if "test_standard" not in st.session_state:
    st.session_state.test_standard = ""

# ==================== 侧边栏配置 ====================
with st.sidebar:
    st.image("https://img.icons8.com/color/96/000000/test-passed.png", width=80)
    st.markdown("## 配置面板")

    with st.expander("🏢 公司信息", expanded=False):
        company_name = st.text_input("公司名称", value=st.session_state.company_name)
        report_title = st.text_input("报告标题", value=st.session_state.report_title)
        tester_name = st.text_input("测试人员", value=st.session_state.tester_name)
        client_name = st.text_input("客户名称", value=st.session_state.client_name)
        test_standard = st.text_input("测试标准", value=st.session_state.test_standard)
        if company_name != st.session_state.company_name:
            st.session_state.company_name = company_name
        if report_title != st.session_state.report_title:
            st.session_state.report_title = report_title
        if tester_name != st.session_state.tester_name:
            st.session_state.tester_name = tester_name
        if client_name != st.session_state.client_name:
            st.session_state.client_name = client_name
        if test_standard != st.session_state.test_standard:
            st.session_state.test_standard = test_standard

    uploaded_file = st.file_uploader(
        "📂 上传Excel文件（支持原始测试数据或编辑后导出的文件）",
        type=["xls", "xlsx"],
        key="file_uploader"
    )
    st.markdown("---")

# ==================== 主区域：根据数据源处理 ====================
if uploaded_file is None:
    st.info("👈 请从左侧侧边栏上传Excel文件")
    st.stop()

# 检测文件是否切换
_file_key = (uploaded_file.name, uploaded_file.size)
if "_last_file_key" in st.session_state and st.session_state._last_file_key != _file_key:
    for k in ["test_ids", "raw_group_data", "edited_data", "all_props",
              "batch_no", "timestamp", "test_type", "selected_groups"]:
        if k in st.session_state:
            if k == "test_ids":
                st.session_state[k] = []
            elif k == "raw_group_data":
                st.session_state[k] = None
            elif k == "edited_data":
                st.session_state[k] = {}
            elif k == "selected_groups":
                st.session_state[k] = set()
            elif k == "all_props":
                st.session_state[k] = None
            elif k in ("batch_no", "timestamp", "test_type"):
                if k == "test_type":
                    st.session_state[k] = "拉伸性能测试"
                else:
                    st.session_state[k] = ""
st.session_state._last_file_key = _file_key

# 检测文件类型并加载数据（仅在 session_state 中尚无数据时解析）
if not st.session_state.test_ids:
    try:
        xlsx = pd.ExcelFile(uploaded_file)
        is_exported = "元数据" in xlsx.sheet_names
    except Exception:
        is_exported = False

    if is_exported:
        if import_edited_data(uploaded_file):
            st.success("成功导入编辑后的数据文件")
        else:
            st.error("导入失败，请检查文件格式")
            st.stop()
    else:
        df_raw = pd.read_excel(uploaded_file, sheet_name=0, header=None)
        groups = find_all_test_groups(df_raw)
        if not groups:
            st.error("❌ 未找到测试编号，请检查文件格式！")
            st.stop()

        test_ids = []
        all_props_initial = []
        raw_group_data = {}
        col_meta_list = None
        detected_test_type = "拉伸性能测试"  # 默认

        for offset, full_text in groups:
            clean_id, batch_no, timestamp = parse_test_id(full_text)
            dim_info = extract_test_dimensions(df_raw, offset)

            try:
                gauge_length = float(dim_info.get("标距", 50))
            except:
                gauge_length = 50.0
            try:
                area = float(dim_info.get("面积", 0))
            except:
                area = 0.0
            try:
                width = float(dim_info.get("宽度{对边长}", 0))
            except:
                width = 0.0

            all_data, force_kgf, disp, stress, strain = extract_group_data(df_raw, offset)
            if len(force_kgf) == 0:
                continue
            if col_meta_list is None:
                col_meta_list = read_columns_meta(df_raw, offset)

            if col_meta_list:
                has_stress = any("应力" in c for c in col_meta_list)
                has_strain = any("应变" in c for c in col_meta_list)
                if has_stress and has_strain:
                    detected_test_type = "拉伸性能测试"
                else:
                    detected_test_type = "背胶耐候性测试"

            if detected_test_type == "拉伸性能测试":
                props = calculate_mechanical_properties(force_kgf, disp, stress, strain, gauge_length, area)
                raw_group_data[clean_id] = {
                    "force": force_kgf, "disp": disp, "stress": stress, "strain": strain,
                    "gauge_length": gauge_length, "area": area, "width": width, "thickness": 0.0,
                    "_col_meta": col_meta_list, "_all_data": all_data, "dim_info": dim_info,
                }
            else:
                props = calculate_peel_properties(force_kgf, disp, width)
                raw_group_data[clean_id] = {
                    "force": force_kgf, "disp": disp, "stress": None, "strain": None,
                    "gauge_length": gauge_length, "area": area, "width": width, "thickness": 0.0,
                    "_col_meta": col_meta_list, "_all_data": all_data, "dim_info": dim_info,
                }
            test_ids.append(clean_id)
            all_props_initial.append(props)

        if not test_ids:
            st.error("无有效数据组")
            st.stop()

        st.session_state.test_type = detected_test_type

        def _init_xy_defaults(first_test_id):
            cols = list(st.session_state.edited_data[first_test_id].columns)
            x_idx = y_idx = None
            for i, n in enumerate(cols):
                if "应变" in n:
                    x_idx = i
                if "应力" in n:
                    y_idx = i
            if x_idx is None or y_idx is None:
                for i, n in enumerate(cols):
                    if "位移" in n:
                        x_idx = i
                    if "荷重" in n or ("力" in n and "应" not in n):
                        y_idx = i
            if x_idx is None:
                x_idx = 0
            if y_idx is None:
                y_idx = min(1, len(cols) - 1)
            if x_idx == y_idx:
                y_idx = (x_idx + 1) % len(cols)
            st.session_state.x_var = cols[x_idx]
            st.session_state.y_var = cols[y_idx]
            st.session_state.x_label = cols[x_idx]
            st.session_state.y_label = cols[y_idx]

        st.session_state.raw_group_data = raw_group_data
        st.session_state.all_props = all_props_initial
        st.session_state.batch_no = batch_no
        st.session_state.timestamp = timestamp

        for tid in test_ids:
            if tid not in st.session_state.edited_data:
                raw = raw_group_data[tid]
                col_meta = raw["_col_meta"]
                all_cols = raw["_all_data"]
                target_len = len(raw["force"])
                df_dict = {}
                for idx, col_name in enumerate(col_meta):
                    arr = all_cols.get(idx, [])
                    if len(arr) >= target_len:
                        df_dict[col_name] = arr[:target_len]
                    elif len(arr) > 0:
                        try:
                            padded = np.pad(arr.astype(float), (0, target_len - len(arr)), constant_values=np.nan)
                            df_dict[col_name] = padded
                        except:
                            padded = list(arr) + [""] * (target_len - len(arr))
                            df_dict[col_name] = padded
                    else:
                        df_dict[col_name] = [np.nan] * target_len
                df = pd.DataFrame(df_dict).reset_index(drop=True)
                st.session_state.edited_data[tid] = df
        st.session_state.test_ids = test_ids

        if st.session_state.test_ids:
            _init_xy_defaults(st.session_state.test_ids[0])

# ==================== 显示概览指标 ====================
col1, col2, col3 = st.columns(3)
col1.metric("测试批号", st.session_state.batch_no)
col2.metric("测试组数", len(st.session_state.test_ids))
col3.metric("测试日期", st.session_state.timestamp)

with st.expander("📋 测试组详细信息", expanded=False):
    first_tid = st.session_state.test_ids[0]
    first_dim = st.session_state.raw_group_data[first_tid].get("dim_info", {})
    meta_parts = []
    meta_mappings = [
        ("试品名称", "试品名称", ""),
        ("长度单位", "长度单位", "mm"),
        ("荷重单位", "荷重单位", "N"),
        ("形状", "形状", ""),
    ]
    for display_label, key, default_val in meta_mappings:
        value = first_dim.get(key, default_val)
        if value:
            meta_parts.append(f"**{display_label}**: {value}")
    if meta_parts:
        st.markdown("　".join(meta_parts))
    elif first_dim:
        st.caption(f"解析到的维度键名：{', '.join(first_dim.keys())}")

    common_keys = ["面积", "宽度{对边长}", "厚度", "倒角半径", "标距", "长度"]
    dim_data = []
    for tid in st.session_state.test_ids:
        dim_info = st.session_state.raw_group_data[tid].get("dim_info", {})
        entry = {"测试编号": tid}
        for k in common_keys:
            entry[k] = dim_info.get(k, "")
        dim_data.append(entry)
    dim_df = pd.DataFrame(dim_data).set_index("测试编号")
    st.dataframe(dim_df, width="stretch")

# ==================== 数据编辑器 ====================
st.subheader("✏️ 数据编辑器")

if "_de_bump" not in st.session_state:
    st.session_state._de_bump = 0

with st.expander("点击展开/收起数据编辑器", expanded=False):
    if not st.session_state.test_ids:
        st.info("无测试组数据")
    else:
        edit_group = st.selectbox("选择要编辑的测试组", st.session_state.test_ids, key="de_edit_group")
        full_df = st.session_state.edited_data[edit_group].copy()

        with st.expander("🔽 按列筛选", expanded=False):
            filter_col = st.selectbox("筛选列", full_df.columns, key=f"de_fcol_{edit_group}")
            col_data = full_df[filter_col]
            if pd.api.types.is_numeric_dtype(col_data):
                lo, hi = float(col_data.min()), float(col_data.max())
                if lo == hi:
                    st.info(f"列「{filter_col}」所有值均为 {lo}")
                    mask = pd.Series(True, index=full_df.index)
                else:
                    rng = st.slider(f"范围：{filter_col}", lo, hi, (lo, hi), key=f"de_fsl_{edit_group}")
                    mask = (full_df[filter_col] >= rng[0]) & (full_df[filter_col] <= rng[1])
            else:
                vals = col_data.dropna().unique().tolist()
                sel = st.multiselect(f"值：{filter_col}", vals, default=vals, key=f"de_fms_{edit_group}")
                mask = full_df[filter_col].isin(sel) if sel else pd.Series(False, index=full_df.index)

            filtered_indices = full_df.index[mask].tolist()
            display_df = full_df.loc[mask].reset_index(drop=True)

        st.caption(f"当前显示 {len(display_df)} / {len(full_df)} 行"
                   + (" · ⚠️ 筛选视图" if len(display_df) != len(full_df) else ""))

        de_key = f"de_editor_{edit_group}_{st.session_state._de_bump}"
        edited_df = st.data_editor(
            display_df,
            key=de_key,
            num_rows="dynamic",
            width="stretch",
            hide_index=True,
        )

        col_btn1, col_btn2, _ = st.columns([2, 2, 4])
        with col_btn1:
            save_clicked = st.button("✅ 保存修改", key="de_save")
        with col_btn2:
            reset_clicked = st.button("🔄 重置此组数据", key="de_reset")

        if save_clicked:
            other_indices = [i for i in range(len(full_df)) if i not in filtered_indices]
            other_rows = full_df.iloc[other_indices].reset_index(drop=True)
            new_full = pd.concat([edited_df, other_rows], ignore_index=True)
            new_full = new_full[full_df.columns]
            st.session_state.edited_data[edit_group] = new_full
            st.session_state._de_bump += 1
            recalc_all_properties()
            st.success("已保存修改")
            st.rerun()

        if reset_clicked:
            reset_data_for_group(edit_group)
            st.success("已重置为原始数据")
            st.rerun()

# ==================== 图表定制与预览 ====================
st.markdown("---")
st.subheader("🎨 图表定制与预览")

all_ids = st.session_state.test_ids
if st.session_state.get("last_ids") != all_ids:
    st.session_state.selected_groups = set(all_ids)
    st.session_state.last_ids = list(all_ids)

st.markdown("##### 测试组")
n_ids = len(all_ids)
group_cols = st.columns(min(8, n_ids))
for i, tid in enumerate(all_ids):
    with group_cols[i % len(group_cols)]:
        checked = st.checkbox(tid, value=tid in st.session_state.selected_groups, key=f"chk_{tid}")
        if checked:
            st.session_state.selected_groups.add(tid)
        else:
            st.session_state.selected_groups.discard(tid)

selected_groups = [tid for tid in all_ids if tid in st.session_state.selected_groups]
if not selected_groups:
    st.warning("请至少选择一个测试组")
    st.stop()

all_cols = list(st.session_state.edited_data[st.session_state.test_ids[0]].columns)

if st.session_state.x_var not in all_cols:
    for c in all_cols:
        if "位移" in c:
            st.session_state.x_var = c
            break
    else:
        st.session_state.x_var = all_cols[0]
    st.session_state.x_label = st.session_state.x_var
if st.session_state.y_var not in all_cols:
    for c in all_cols:
        if "荷重" in c or ("力" in c and "应" not in c):
            st.session_state.y_var = c
            break
    else:
        st.session_state.y_var = all_cols[1] if len(all_cols) > 1 else all_cols[0]
    st.session_state.y_label = st.session_state.y_var

def _sync_x_label():
    st.session_state.x_label = st.session_state.x_var
def _sync_y_label():
    st.session_state.y_label = st.session_state.y_var

with st.container(border=True):
    st.markdown("**坐标轴**")
    r1 = st.columns(5)
    with r1[0]:
        st.selectbox("X 轴变量", all_cols, key="x_var", on_change=_sync_x_label)
    with r1[1]:
        st.text_input("X 轴标签", key="x_label")
    with r1[3]:
        st.selectbox("Y 轴变量", all_cols, key="y_var", on_change=_sync_y_label)
    with r1[4]:
        st.text_input("Y 轴标签", key="y_label")

_CHART_PALETTE = ['#1a237e', '#ef5350', '#2e7d32', '#ff8f00', '#6a1b9a',
                  '#00838f', '#d81b60', '#3e2723', '#558b2f', '#01579b']

with st.container(border=True):
    st.markdown("**样式**")
    sr = st.columns([1, 1, 2])
    with sr[2]:
        st.slider("点大小", 1, 10, value=2, step=1, key="line_width")
    with sr[0]:
        st.radio("颜色模式", ["自动分配", "自定义每个测试组"], key="color_mode", horizontal=True)

    custom_colors = {}
    if st.session_state.color_mode == "自定义每个测试组":
        nc = min(6, len(selected_groups))
        color_cols = st.columns(nc)
        for i, test_id in enumerate(selected_groups):
            with color_cols[i % nc]:
                default = st.session_state.get(
                    f"color_{test_id}", _CHART_PALETTE[i % len(_CHART_PALETTE)]
                )
                custom_colors[test_id] = st.color_picker(test_id, value=default, key=f"color_{test_id}")

x_var = st.session_state.x_var
y_var = st.session_state.y_var
x_label = st.session_state.x_label or x_var
y_label = st.session_state.y_label or y_var
line_width = st.session_state.get("line_width", 2)
color_mode = st.session_state.color_mode

# 对选中的测试组按原始顺序排序，保证图例顺序一致
selected_groups_sorted = sorted(selected_groups, key=lambda x: st.session_state.test_ids.index(x))

fig = plot_plotly_chart(
    selected_groups_sorted, x_var, y_var, "无筛选",
    0.0, 100.0, 0, 0,
    x_label, y_label, line_width, color_mode, custom_colors,
    st.session_state.edited_data
)
st.plotly_chart(fig, width='stretch', config={'displayModeBar': True, 'responsive': True})

# ==================== 导出及报告按钮 ====================
st.markdown("---")

# 提取所有必要的 session_state 数据为局部变量，避免在 lambda 中直接访问
company_name = st.session_state.company_name
report_title = st.session_state.report_title
batch_no = st.session_state.batch_no
timestamp = st.session_state.timestamp
tester_name = st.session_state.tester_name
client_name = st.session_state.client_name
test_standard = st.session_state.test_standard
test_type = st.session_state.test_type
all_props = st.session_state.all_props
raw_group_data = st.session_state.raw_group_data
test_ids = st.session_state.test_ids
edited_data = st.session_state.edited_data
x_var = st.session_state.x_var
y_var = st.session_state.y_var
x_label = st.session_state.x_label
y_label = st.session_state.y_label
line_width = st.session_state.get("line_width", 2)
color_mode = st.session_state.color_mode

# Excel 导出
col_ex, _ = st.columns([1, 2])
with col_ex:
    dl_excel = st.download_button(
        key="dl_excel",
        label="📥 导出当前所有数据为Excel",
        data=lambda: export_edited_data(
            test_ids, raw_group_data, all_props, edited_data
        ).getvalue(),
        file_name=f"{batch_no}_编辑后数据.xls",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        width="stretch",
    )

st.markdown("---")
st.subheader("📄 生成报告")

# 收集自定义颜色
_custom_colors = {}
if st.session_state.color_mode == "自定义每个测试组":
    for tid in selected_groups:
        v = st.session_state.get(f"color_{tid}")
        if v:
            _custom_colors[tid] = v

# 自定义颜色元组
custom_colors_tuple = tuple(sorted(
    [(tid, st.session_state.get(f"color_{tid}", "")) for tid in selected_groups_sorted if st.session_state.get(f"color_{tid}")]
))

col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
with col_btn2:
    dl_word = st.download_button(
        label="📄 下载Word报告",
        data=lambda: generate_word_report_bytes(
            selected_groups_sorted,
            x_var, y_var, x_label, y_label,
            line_width, color_mode, custom_colors_tuple,
            company_name, report_title, batch_no, timestamp,
            tester_name, client_name, test_standard,
            test_type, all_props, raw_group_data, test_ids,
            edited_data
        ),
        file_name=f"{batch_no}_测试报告.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        width="stretch",
    )

    dl_html = st.download_button(
        key="dl_html",
        label="🌐 下载可交互HTML报告",
        data=lambda: generate_html_report(
            test_ids, all_props, batch_no, timestamp,
            x_var, y_var, x_label, y_label, "无筛选",
            0.0, 100.0, 0, 0,
            selected_groups_sorted,
            line_width, _custom_colors,
            tester_name, client_name, test_standard,
            test_type, edited_data, raw_group_data,
            company_name, report_title
        ),
        file_name=f"{batch_no}_测试报告.html",
        mime="text/html",
        width="stretch",
    )
